"""Homebrew Workshop routes.

Rook can read a fillable Markdown/Text/Word template, turn it into structured
homebrew, let the user review the draft, then save it into the user's personal
or campaign-scoped homebrew library.
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Response
from pydantic import BaseModel
from typing import Optional, Dict, Any, List
import io
import json
import re
import uuid
from datetime import datetime, timezone

from utils.auth import get_current_user
from config import db, logger
from utils.llm_provider import LlmChat, UserMessage, get_llm_api_key

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional local/dev dependency guard
    Document = None

router = APIRouter()

CONTENT_TYPES = {
    "race",
    "class",
    "subclass",
    "feat",
    "spell",
    "background",
    "magic_item",
    "monster",
    "npc",
    "custom_rule",
}

CONTENT_TYPE_ALIASES = {
    "species": "race",
    "ancestry": "race",
    "origin": "race",
    "item": "magic_item",
    "magicitem": "magic_item",
    "magic-item": "magic_item",
    "magic item": "magic_item",
    "rule": "custom_rule",
    "customrule": "custom_rule",
    "custom-rule": "custom_rule",
    "custom rule": "custom_rule",
    "sub-class": "subclass",
    "sub class": "subclass",
}

COLLECTION = {
    "race": "user_races",
    "class": "user_classes",
    "subclass": "user_subclasses",
    "feat": "user_feats",
    "spell": "user_spells",
    "background": "user_backgrounds",
    "magic_item": "user_magic_items",
    "monster": "user_monsters",
    "npc": "user_npcs",
    "custom_rule": "user_custom_rules",
}

ADVANCED_MECHANIC_HINTS = {
    "resources": "[{name, formula, max, regain, spend_triggers, visible_on_sheet}] — custom pools such as Scarab Charges = warlock level or Greed Tokens = proficiency bonus",
    "actions": "[{name, action_type, cost, resource_cost, description}] — sheet actions and spendable options",
    "passive_effects": "[{name, target, mode, value, condition, formula}] — bonuses, tags, speeds, proficiencies, resistances, etc.",
    "scaling": "[{level, formula, description}] — level, proficiency, ability, class-level, or spell-level scaling",
    "upgrades": "[{level, name, description, replaces_or_improves}] — feature improvements over time",
    "automation_notes": "string — how Rook/the sheet should wire this into builders and character sheets",
}

SCHEMA_HINTS = {
    "race": {
        "name": "string",
        "description": "string",
        "size": "Tiny|Small|Medium|Large",
        "speed": "int (default 30)",
        "ability_bonuses": "object e.g. {strength: 1, dexterity: 2}",
        "traits": "[{name, description, level, rules_text}]",
        "languages": "[string]",
        "subraces": "[{name, description, ability_bonuses, traits}]",
        **ADVANCED_MECHANIC_HINTS,
    },
    "class": {
        "name": "string",
        "description": "string",
        "hit_die": "d6|d8|d10|d12",
        "primary_ability": "string",
        "saving_throw_proficiencies": "[string] — two abilities",
        "armor_proficiencies": "[string]",
        "weapon_proficiencies": "[string]",
        "tool_proficiencies": "[string]",
        "skill_choices": "object e.g. {choose: 2, from: [athletics, insight]}",
        "equipment": "[string]",
        "features": "[{level: int 1-20, name, description, rules_text, resources, actions, passive_effects}]",
        "subclass_unlock_levels": "[int] — levels where subclass features are gained",
        **ADVANCED_MECHANIC_HINTS,
    },
    "subclass": {
        "name": "string",
        "parent_class": "string — must match a class name",
        "description": "string",
        "subclass_level": "int (level the subclass unlocks, usually 3)",
        "features": "[{level: int 1-20, name, description, rules_text, resources, actions, passive_effects, upgrades}]",
        **ADVANCED_MECHANIC_HINTS,
    },
    "feat": {
        "name": "string",
        "description": "string",
        "prerequisite": "string",
        "repeatable": "bool",
        "ability_score_increase": "object e.g. {choose: 1, from: [strength, dexterity], amount: 1}",
        "benefits": "[{name, description, rules_text}]",
        **ADVANCED_MECHANIC_HINTS,
    },
    "spell": {
        "name": "string",
        "description": "string",
        "level": "int 0-9",
        "school": "string",
        "casting_time": "string",
        "range": "string",
        "components": "string or [string]",
        "duration": "string",
        "ritual": "bool",
        "concentration": "bool",
        "classes": "[string]",
        "damage": "object e.g. {dice: 2d6, type: fire}",
        "higher_level": "string",
        "effects": "[string]",
        **ADVANCED_MECHANIC_HINTS,
    },
    "background": {
        "name": "string",
        "description": "string",
        "skill_proficiencies": "[string] — usually two skills",
        "tool_proficiencies": "[string]",
        "languages": "int — number of additional languages",
        "equipment": "[string]",
        "feature_name": "string",
        "feature_description": "string",
        "suggested_characteristics": "object or [string]",
        **ADVANCED_MECHANIC_HINTS,
    },
    "magic_item": {
        "name": "string",
        "type": "Weapon|Armor|Wondrous Item|Potion|Ring|Rod|Scroll|Staff|Wand",
        "rarity": "common|uncommon|rare|very rare|legendary|artifact",
        "requires_attunement": "bool",
        "attunement_requirement": "string",
        "description": "string",
        "effects": "[string] — short bullet effects",
        "charges": "object e.g. {max: 3, regain: 1d3 at dawn}",
        **ADVANCED_MECHANIC_HINTS,
    },
    "monster": {
        "name": "string",
        "creature_type": "string, e.g. undead, fiend, beast, construct",
        "size": "Tiny|Small|Medium|Large|Huge|Gargantuan",
        "alignment": "string",
        "challenge_rating": "string or number",
        "armor_class": "int",
        "hit_points": "int",
        "speed": "string, e.g. 30 ft., fly 60 ft.",
        "abilities": "object with str,dex,con,int,wis,cha numbers",
        "saving_throws": "object or [string]",
        "skills": "object or [string]",
        "damage_resistances": "[string]",
        "damage_immunities": "[string]",
        "condition_immunities": "[string]",
        "senses": "string",
        "languages": "string",
        "traits": "[{name, description}]",
        "actions": "[{name, description, attack_bonus, damage}]",
        "bonus_actions": "[{name, description}]",
        "reactions": "[{name, description}]",
        "legendary_actions": "[{name, description}]",
        "lair_actions": "[{name, description}]",
        "description": "string",
        "role": "solo|brute|skirmisher|controller|minion|boss|support",
        **ADVANCED_MECHANIC_HINTS,
    },
    "npc": {
        "name": "string",
        "ancestry": "string",
        "role": "string, e.g. innkeeper, rival, captain, patron",
        "faction": "string",
        "location": "string",
        "appearance": "string",
        "personality": "string",
        "voice": "string",
        "mannerisms": "[string]",
        "ideal": "string",
        "bond": "string",
        "flaw": "string",
        "secret": "string",
        "wants": "string",
        "fears": "string",
        "quest_hooks": "[string]",
        "combat_role": "noncombatant|minion|rival|ally|boss|support",
        "stat_hint": "string, e.g. commoner, bandit captain, mage",
        "description": "string",
        **ADVANCED_MECHANIC_HINTS,
    },
    "custom_rule": {
        "name": "string",
        "category": "exploding_dice|chaos_tokens|custom_skill|resting|combat|magic|futuristic|other",
        "summary": "string",
        "enabled_by_default": "bool",
        "rule_text": "string",
        "trigger": "string",
        "resolution": "string",
        "examples": "[string]",
        "settings": "object. For exploding dice include dice: [d4,d6,d8,d10,d12,d20]. For skills include ability and description.",
        "player_visible": "bool",
        "gm_notes": "string",
        **ADVANCED_MECHANIC_HINTS,
    },
}

TEMPLATE_LABELS = {
    "race": "Race / Species",
    "class": "Class",
    "subclass": "Subclass",
    "feat": "Feat",
    "spell": "Spell",
    "background": "Background",
    "magic_item": "Magic Item",
    "monster": "Monster / Creature",
    "npc": "NPC",
    "custom_rule": "Custom Rule",
}

FIELD_PROMPTS = {
    "name": "The public name shown in builders, sheets, and libraries.",
    "description": "The plain-English theme, lore, table-facing summary, and what makes it different.",
    "parent_class": "For subclasses only. Example: Warlock, Fighter, Monk.",
    "features": "Feature blocks. Include level, name, rules text, resource costs, and upgrades.",
    "resources": "Custom pools/charges/tokens. Include formula, max, regain, spending rules, and whether it appears on the sheet.",
    "actions": "Buttons/options the sheet should show. Include action type and resource cost.",
    "passive_effects": "Bonuses, tags, resistances, speed changes, or proficiencies that should appear or apply on the sheet.",
    "scaling": "How numbers scale by class level, character level, proficiency bonus, ability modifier, or spell level.",
    "upgrades": "Later feature improvements and what they replace or improve.",
    "automation_notes": "Tell Rook exactly how this should become usable on the site.",
}

EXAMPLE_SNIPPETS = {
    "subclass": """## Name
The Gilded Scarab

## Parent Class
Warlock

## Subclass Level
1

## Description
A pact with a glittering tomb-scarab spirit that rewards greed, bargains, and cursed treasure.

## Resources
- Scarab Charges: max = Warlock level, regain all on long rest, visible on sheet, spent by subclass features.
- Greed Tokens: max = Proficiency Bonus, GM adjustable, visible on sheet.

## Features
### Level 1 - Gilded Pact
You gain Scarab Charges equal to your warlock level. You can spend 1 charge when you hit with Eldritch Blast to add necrotic or radiant damage equal to your proficiency bonus.

## Actions
- Spend Scarab Charge: bonus action, cost 1 Scarab Charge, trigger subclass feature.

## Automation Notes
Create sheet-visible resources, spending buttons, and level-scaling max values.
""",
    "feat": """## Name
Shield-Breaker

## Prerequisite
Strength 13 or higher

## Description
You have learned how to punish guarded enemies.

## Benefits
- When you hit a creature carrying a shield, you can push it 5 feet once per turn.

## Passive Effects
- Add +1 Strength, up to a maximum of 20.
""",
    "spell": """## Name
Ashen Rebuke

## Level
1

## School
Evocation

## Casting Time
1 reaction, which you take when a creature damages you

## Range
60 feet

## Components
V, S

## Duration
Instantaneous

## Classes
Warlock, Sorcerer

## Description
Flame and ash snap back at the attacker.

## Damage
2d6 fire

## Higher Level
Increase the damage by 1d6 for each slot level above 1st.
""",
}


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _normalise_content_type(content_type: str) -> str:
    raw = str(content_type or "").strip().lower().replace("-", "_").replace(" ", "_")
    raw = CONTENT_TYPE_ALIASES.get(raw, raw)
    if raw not in CONTENT_TYPES:
        raise HTTPException(status_code=400, detail=f"content_type must be one of {sorted(CONTENT_TYPES)}")
    return raw


def _normalise_edition(edition: str) -> str:
    return "2024" if str(edition or "").strip() == "2024" else "2014"


def _normalise_parsed(content_type: str, parsed: Dict[str, Any], edition: str) -> Dict[str, Any]:
    data = dict(parsed or {})
    if content_type == "subclass" and not data.get("parent_class"):
        data["parent_class"] = data.get("baseClass") or data.get("base_class") or data.get("class") or ""
    if content_type == "magic_item" and data.get("item_type") and not data.get("type"):
        data["type"] = data.get("item_type")
    data["content_type"] = content_type
    data["edition"] = _normalise_edition(edition)
    return data


def _build_template(content_type: str, edition: str = "2014") -> str:
    content_type = _normalise_content_type(content_type)
    edition = _normalise_edition(edition)
    label = TEMPLATE_LABELS[content_type]
    schema = SCHEMA_HINTS[content_type]
    lines = [
        f"# Rook Homebrew Template: {label}",
        "",
        f"Edition: {edition}",
        "",
        "Fill in what you know. Leave anything unknown blank and Rook can help complete it.",
        "Use plain English for rules text. For features, resources, actions, and scaling, be as explicit as possible.",
        "",
    ]
    for field, hint in schema.items():
        title = field.replace("_", " ").title()
        lines.extend([
            f"## {title}",
            FIELD_PROMPTS.get(field, str(hint)),
            "",
        ])
    example = EXAMPLE_SNIPPETS.get(content_type)
    if example:
        lines.extend(["---", "", "# Example", "", example.strip(), ""])
    return "\n".join(lines).strip() + "\n"


def _docx_to_text(file_bytes: bytes) -> str:
    if Document is None:
        raise HTTPException(status_code=503, detail="python-docx not installed on server")
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read .docx file: {type(e).__name__}")
    parts: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))
    return "\n".join(parts)


def _extract_json(reply: str) -> Optional[Dict[str, Any]]:
    if not reply:
        return None
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", reply, re.DOTALL)
    candidate = fenced.group(1) if fenced else None
    if candidate is None:
        first = reply.find("{")
        last = reply.rfind("}")
        if first != -1 and last > first:
            candidate = reply[first:last + 1]
    if candidate is None:
        return None
    try:
        parsed = json.loads(candidate)
        return parsed if isinstance(parsed, dict) else None
    except json.JSONDecodeError:
        logger.warning("Rook returned invalid homebrew JSON")
        return None


def _required_fields_for(content_type: str) -> List[str]:
    return {
        "race": ["name", "size", "speed"],
        "class": ["name", "hit_die", "features"],
        "subclass": ["name", "parent_class", "features"],
        "feat": ["name", "description", "benefits"],
        "spell": ["name", "level", "school", "casting_time"],
        "background": ["name", "skill_proficiencies"],
        "magic_item": ["name", "rarity"],
        "monster": ["name", "armor_class", "hit_points", "challenge_rating", "actions"],
        "npc": ["name", "role", "personality", "secret"],
        "custom_rule": ["name", "category", "summary", "rule_text"],
    }[content_type]


def _flag_missing(content_type: str, parsed: Dict[str, Any]) -> List[str]:
    missing: List[str] = []
    for field in _required_fields_for(content_type):
        value = parsed.get(field)
        if value in (None, "", [], {}):
            missing.append(field)
    return missing


def _merge_draft(original: Dict[str, Any], completed: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(original or {})
    for key, value in (completed or {}).items():
        if value in (None, "", [], {}):
            continue
        merged[key] = value
    return merged


async def _llm_json(content_type: str, system: str, prompt: str, username: str) -> Dict[str, Any]:
    if not LlmChat or not UserMessage or not get_llm_api_key("anthropic"):
        raise HTTPException(status_code=503, detail="Rook is not configured on this server.")
    chat = LlmChat(
        api_key=get_llm_api_key("anthropic"),
        session_id=f"homebrew-{content_type}-{username}-{uuid.uuid4().hex[:6]}",
        system_message=system
    ).with_model("anthropic", "claude-sonnet-4-5-20250929")
    try:
        reply = await chat.send_message(UserMessage(text=prompt))
    except Exception as e:
        logger.exception("Homebrew Rook request failed")
        raise HTTPException(status_code=502, detail=f"Rook request failed: {type(e).__name__}")
    parsed = _extract_json(reply or "")
    return parsed or {}


async def _llm_extract(content_type: str, raw_text: str, username: str, edition: str = "2014") -> Dict[str, Any]:
    schema = SCHEMA_HINTS.get(content_type)
    if not schema:
        raise HTTPException(status_code=400, detail=f"Unsupported content_type '{content_type}'")
    schema_str = json.dumps(schema, indent=2)
    snippet = raw_text[:12000]
    system = (
        "You are Rook, an SRD-compliant TTRPG homebrew extraction assistant. Read the user's homebrew text and "
        "return ONLY a single JSON object that matches the requested schema. Use null, empty arrays, or empty strings "
        "for fields you cannot find. Do not include any explanation. Preserve mechanics such as resources, actions, "
        "passive effects, scaling, upgrades, and automation notes whenever present."
    )
    prompt = (
        f"Content type: {content_type}\n"
        f"Edition: {edition}\n"
        f"Schema (return JSON matching this shape):\n{schema_str}\n\n"
        f"Source text:\n{snippet}\n\n"
        f"Return ONLY the JSON object."
    )
    parsed = await _llm_json(content_type, system, prompt, username)
    return _normalise_parsed(content_type, parsed, edition)


async def _llm_complete(content_type: str, draft: Dict[str, Any], context: str, username: str, edition: str = "2014") -> Dict[str, Any]:
    schema = SCHEMA_HINTS.get(content_type)
    if not schema:
        raise HTTPException(status_code=400, detail=f"Unsupported content_type '{content_type}'")
    schema_str = json.dumps(schema, indent=2)
    draft_str = json.dumps(draft or {}, indent=2)[:12000]
    context_snippet = (context or "")[:6000]
    system = (
        "You are Rook, a practical TTRPG homebrew co-designer. Complete missing or thin fields using the user's existing choices, "
        "theme, tone, and stats. Keep rules usable at the table. Do not overwrite the core identity unless it is clearly empty. "
        "Preserve sheet integration mechanics like resources, actions, passive effects, scaling, upgrades, and automation notes. "
        "Return ONLY one JSON object matching the schema. No markdown."
    )
    prompt = (
        f"Content type: {content_type}\n"
        f"Edition: {edition}\n"
        f"Schema:\n{schema_str}\n\n"
        f"Current draft JSON:\n{draft_str}\n\n"
        f"Optional user source/context:\n{context_snippet}\n\n"
        "Fill missing or weak fields and return the complete JSON object."
    )
    completed = await _llm_json(content_type, system, prompt, username)
    return _normalise_parsed(content_type, _merge_draft(draft or {}, completed), edition)


@router.get("/homebrew/templates")
async def list_templates():
    return {
        "templates": [
            {"content_type": key, "label": TEMPLATE_LABELS[key], "filename": f"rook-homebrew-{key}.md"}
            for key in sorted(CONTENT_TYPES)
        ]
    }


@router.get("/homebrew/template/{content_type}")
async def download_template(content_type: str, edition: str = "2014"):
    normalised_type = _normalise_content_type(content_type)
    normalised_edition = _normalise_edition(edition)
    text = _build_template(normalised_type, normalised_edition)
    filename = f"rook-homebrew-{normalised_type}-{normalised_edition}.md"
    return Response(
        content=text,
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/homebrew/parse-docx")
async def parse_docx(
    content_type: str = Form(...),
    file: UploadFile = File(...),
    edition: str = Form("2014"),
    username: str = Depends(get_current_user)
):
    normalised_type = _normalise_content_type(content_type)
    normalised_edition = _normalise_edition(edition)

    filename = (file.filename or "").lower()
    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=400, detail="Empty file")

    if filename.endswith(".docx"):
        text = _docx_to_text(raw_bytes)
    elif filename.endswith((".txt", ".md", ".markdown")):
        try:
            text = raw_bytes.decode("utf-8", errors="ignore")
        except Exception:
            raise HTTPException(status_code=400, detail="Could not read text file")
    else:
        raise HTTPException(status_code=400, detail="Only .docx, .txt, .md, or .markdown files are supported")

    if not text.strip():
        raise HTTPException(status_code=400, detail="The file appears to be empty.")

    draft = await _llm_extract(normalised_type, text, username, normalised_edition)
    missing = _flag_missing(normalised_type, draft)
    return {
        "content_type": normalised_type,
        "edition": normalised_edition,
        "draft": draft,
        "missing_fields": missing,
        "source_filename": file.filename,
        "source_excerpt": text[:1500],
    }


class ParseTextRequest(BaseModel):
    content_type: str
    edition: str = "2014"
    text: str


@router.post("/homebrew/parse-text")
async def parse_text(req: ParseTextRequest, username: str = Depends(get_current_user)):
    normalised_type = _normalise_content_type(req.content_type)
    normalised_edition = _normalise_edition(req.edition)
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="Text is empty")
    draft = await _llm_extract(normalised_type, req.text, username, normalised_edition)
    missing = _flag_missing(normalised_type, draft)
    return {
        "content_type": normalised_type,
        "edition": normalised_edition,
        "draft": draft,
        "missing_fields": missing,
        "source_excerpt": req.text[:1500],
    }


class CompleteDraftRequest(BaseModel):
    content_type: str
    edition: str = "2014"
    draft: Dict[str, Any]
    context: str = ""


@router.post("/homebrew/complete-draft")
async def complete_draft(req: CompleteDraftRequest, username: str = Depends(get_current_user)):
    normalised_type = _normalise_content_type(req.content_type)
    normalised_edition = _normalise_edition(req.edition)
    completed = await _llm_complete(normalised_type, req.draft or {}, req.context or "", username, normalised_edition)
    return {
        "content_type": normalised_type,
        "edition": normalised_edition,
        "draft": completed,
        "missing_fields": _flag_missing(normalised_type, completed),
    }


class HomebrewSaveRequest(BaseModel):
    content_type: str
    edition: str = "2014"
    data: Dict[str, Any]
    ruleset_id: Optional[str] = None
    homebrew_id: Optional[str] = None
    visibility: str = "private"
    campaign_id: Optional[str] = None


@router.post("/homebrew/save")
async def save_homebrew(req: HomebrewSaveRequest, username: str = Depends(get_current_user)):
    normalised_type = _normalise_content_type(req.content_type)
    normalised_edition = _normalise_edition(req.edition)
    visibility = req.visibility if req.visibility in {"private", "campaign", "shared_copy", "public"} else "private"
    if visibility == "campaign" and not req.campaign_id:
        raise HTTPException(status_code=400, detail="campaign_id is required for campaign homebrew")

    coll_name = COLLECTION[normalised_type]
    ruleset_id = req.ruleset_id
    if not ruleset_id:
        existing = await db.user_rulesets.find_one({"user_id": username, "name": "Homebrew Workshop", "edition": normalised_edition}, {"_id": 0})
        if existing:
            ruleset_id = existing["id"]
        else:
            ruleset_id = str(uuid.uuid4())
            await db.user_rulesets.insert_one({
                "id": ruleset_id,
                "user_id": username,
                "name": "Homebrew Workshop",
                "description": "Rook-assisted homebrew content",
                "edition": normalised_edition,
                "version": "1.0",
                "is_active": True,
                "created_at": _now(),
            })

    doc = _normalise_parsed(normalised_type, dict(req.data or {}), normalised_edition)
    doc.update({
        "user_id": username,
        "owner_user_id": username,
        "created_by_user_id": username,
        "ruleset_id": ruleset_id,
        "edition": normalised_edition,
        "visibility": visibility,
        "campaign_id": req.campaign_id if visibility == "campaign" else req.campaign_id,
        "source": "Homebrew Workshop",
        "source_type": "user_homebrew",
        "license": doc.get("license") or "user_provided_private_use",
        "share_policy": doc.get("share_policy") or {
            "allowPrivateShare": True,
            "allowCampaignUse": True,
            "allowPublicListing": visibility == "public",
        },
        "updated_at": _now(),
    })

    if req.homebrew_id:
        existing = await db[coll_name].find_one({"id": req.homebrew_id, "user_id": username})
        if not existing:
            raise HTTPException(status_code=404, detail="Homebrew item not found")
        await db[coll_name].update_one(
            {"id": req.homebrew_id, "user_id": username},
            {"$set": {k: v for k, v in doc.items() if k != "_id"}},
        )
        doc["id"] = req.homebrew_id
    else:
        doc["id"] = str(uuid.uuid4())
        doc["created_at"] = _now()
        await db[coll_name].insert_one(doc)

    doc.pop("_id", None)
    return {"saved": True, "content_type": normalised_type, "homebrew": doc}


@router.get("/homebrew")
async def list_homebrew(content_type: Optional[str] = None, edition: Optional[str] = None, username: str = Depends(get_current_user)):
    types = [_normalise_content_type(content_type)] if content_type else sorted(CONTENT_TYPES)
    out: Dict[str, List[Dict[str, Any]]] = {}
    for t in types:
        coll_name = COLLECTION.get(t)
        if not coll_name:
            continue
        q: Dict[str, Any] = {"user_id": username}
        if edition:
            q["edition"] = _normalise_edition(edition)
        cursor = db[coll_name].find(q, {"_id": 0})
        out[t] = [item async for item in cursor]
    return {"homebrew": out}


@router.delete("/homebrew/{content_type}/{homebrew_id}")
async def delete_homebrew(content_type: str, homebrew_id: str, username: str = Depends(get_current_user)):
    normalised_type = _normalise_content_type(content_type)
    coll_name = COLLECTION[normalised_type]
    result = await db[coll_name].delete_one({"id": homebrew_id, "user_id": username})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Homebrew item not found")
    return {"deleted": homebrew_id, "content_type": normalised_type}
